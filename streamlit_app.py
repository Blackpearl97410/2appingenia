from pathlib import Path
from io import BytesIO
import re

import pandas as pd
import streamlit as st
from docx import Document
from pypdf import PdfReader
from pypdf.errors import PdfReadError


ROOT_DIR = Path(__file__).parent
SAMPLE_CSV = ROOT_DIR / "data" / "samples" / "converted_data.csv"


@st.cache_data
def load_demo_data() -> pd.DataFrame:
    if not SAMPLE_CSV.exists():
        return pd.DataFrame()
    return pd.read_csv(SAMPLE_CSV)


def docx_to_markdown(document: Document) -> str:
    blocks = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and "heading" in paragraph.style.name.lower():
            blocks.append(f"## {text}")
        else:
            blocks.append(text)
    return "\n\n".join(blocks)


def get_uploaded_suffix(uploaded_file) -> str:
    return Path(uploaded_file.name).suffix.lower() or "inconnu"


def get_uploaded_bytes(uploaded_file) -> bytes:
    return uploaded_file.getvalue()


def parse_text_bytes(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")


def parse_csv_bytes(file_bytes: bytes) -> pd.DataFrame:
    return pd.read_csv(BytesIO(file_bytes))


def parse_excel_bytes(file_bytes: bytes) -> dict[str, pd.DataFrame]:
    return pd.read_excel(BytesIO(file_bytes), sheet_name=None)


def parse_pdf_bytes(file_bytes: bytes) -> tuple[str, int, int, str | None]:
    try:
        reader = PdfReader(BytesIO(file_bytes))
        pages_text = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages_text.append(page_text.strip())
        return "\n\n".join(pages_text), len(reader.pages), len(pages_text), None
    except PdfReadError as exc:
        return "", 0, 0, f"PDF illisible par le parseur ({exc})"
    except Exception as exc:
        return "", 0, 0, f"Erreur de lecture PDF ({exc.__class__.__name__})"


def parse_docx_bytes(file_bytes: bytes) -> tuple[str, str, int]:
    document = Document(BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    text_content = "\n\n".join(paragraphs)
    markdown_content = docx_to_markdown(document)
    return text_content, markdown_content, len(paragraphs)


def extract_text_metadata(text: str, filename: str) -> dict[str, str]:
    compact_text = re.sub(r"\s+", " ", text).strip()
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    title = lines[0] if lines else Path(filename).stem

    date_match = re.search(
        r"\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}-\d{2}-\d{2})\b",
        compact_text,
    )
    amount_match = re.search(
        r"\b\d[\d\s.,]{2,}\s?(?:€|euros?)\b",
        compact_text,
        flags=re.IGNORECASE,
    )
    org_match = re.search(
        r"\b(Région|Region|CNM|ADEME|BPI|France Travail|DEETS|Europe|Etat|Minist[eè]re)\b",
        compact_text,
        flags=re.IGNORECASE,
    )

    lower_name = filename.lower()
    if "appel" in lower_name or "aap" in lower_name:
        document_type = "appel a projets"
    elif "formulaire" in lower_name:
        document_type = "formulaire"
    elif "cadre" in lower_name:
        document_type = "cadre d'intervention"
    elif "reglement" in lower_name:
        document_type = "reglement"
    else:
        document_type = "document texte"

    return {
        "Titre probable": title[:120],
        "Type probable": document_type,
        "Date detectee": date_match.group(1) if date_match else "Non detectee",
        "Montant detecte": amount_match.group(0) if amount_match else "Non detecte",
        "Organisme detecte": org_match.group(1) if org_match else "Non detecte",
    }


def extract_table_metadata(dataframe: pd.DataFrame, filename: str) -> dict[str, str]:
    return {
        "Nom du fichier": filename,
        "Type probable": "tableau",
        "Nombre de lignes": str(len(dataframe)),
        "Nombre de colonnes": str(len(dataframe.columns)),
        "Colonnes detectees": ", ".join(str(col) for col in list(dataframe.columns)[:6]) or "Aucune",
    }


def render_metadata(metadata: dict[str, str]) -> None:
    st.markdown("### Metadonnees detectees")
    cols = st.columns(2)
    items = list(metadata.items())
    for index, (label, value) in enumerate(items):
        cols[index % 2].write(f"**{label}** : {value}")


def normalize_detected_value(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def add_detected_value(store: dict[str, set[str]], key: str, value: str, source_name: str) -> None:
    normalized = normalize_detected_value(value)
    if not normalized or normalized in {"Non detectee", "Non detecte", "Aucun", "Aucune"}:
        return
    store.setdefault(normalized, set()).add(source_name)


def format_detected_values(store: dict[str, set[str]]) -> str:
    if not store:
        return "Aucun"
    items = sorted(
        store.items(),
        key=lambda item: (-len(item[1]), item[0].lower()),
    )
    return " | ".join(
        f"{value} ({len(sources)} doc)" if len(sources) == 1 else f"{value} ({len(sources)} docs)"
        for value, sources in items[:8]
    )


def choose_priority_value(block_candidates: dict[str, str], priority_order: list[str]) -> str:
    for block_name in priority_order:
        value = block_candidates.get(block_name)
        if value and value not in {"Aucun", "Aucune"}:
            return f"{value} ({block_name})"
    return "Aucun"


def compute_global_prescore(
    available_blocks: list[str],
    block_criteria_scores: dict[str, int],
    issues: list[str],
) -> tuple[str, str]:
    score = 0

    score += len(available_blocks) * 10
    score += sum(min(points, 30) for points in block_criteria_scores.values())
    score -= min(len(issues), 6) * 10

    score = max(0, min(score, 100))

    if score >= 75:
        label = "bon"
    elif score >= 45:
        label = "moyen"
    else:
        label = "faible"

    return label, f"{score}/100"


def extract_keywords_from_text(text: str) -> list[str]:
    words = re.findall(r"\b[a-zA-ZÀ-ÿ][a-zA-ZÀ-ÿ-]{3,}\b", text.lower())
    filtered = []
    seen = set()
    stop_words = {
        "dans", "avec", "pour", "cette", "document", "documents", "projet",
        "dossier", "client", "bloc", "fichier", "charge", "titre", "type",
        "date", "montant", "organisme", "texte", "tableau", "feuille",
    }
    for word in words:
        if len(word) <= 4 or word in stop_words or word in seen:
            continue
        seen.add(word)
        filtered.append(word)
        if len(filtered) >= 20:
            break
    return filtered


def aggregate_block_text(uploaded_files) -> str:
    chunks = []
    for uploaded_file in uploaded_files:
        suffix = get_uploaded_suffix(uploaded_file)
        file_bytes = get_uploaded_bytes(uploaded_file)

        if suffix in {".txt", ".md"}:
            chunks.append(parse_text_bytes(file_bytes))
        elif suffix == ".docx":
            text_content, _, _ = parse_docx_bytes(file_bytes)
            chunks.append(text_content)
        elif suffix == ".pdf":
            pdf_text, _, _, _ = parse_pdf_bytes(file_bytes)
            chunks.append(pdf_text)
        elif suffix == ".csv":
            dataframe = parse_csv_bytes(file_bytes)
            chunks.append(" ".join(str(col) for col in dataframe.columns))
        elif suffix == ".xlsx":
            workbook = parse_excel_bytes(file_bytes)
            business_sheets, informative_sheets = filter_business_sheets(workbook)
            for sheet_name, sheet_df in business_sheets.items():
                chunks.append(sheet_name)
                chunks.append(" ".join(str(col) for col in sheet_df.columns))
            for sheet_name in informative_sheets:
                chunks.append(sheet_name)

    return "\n".join(chunk for chunk in chunks if chunk).lower()


def evaluate_block_criteria(block_name: str, uploaded_files, insights: dict[str, str]) -> dict[str, str | int]:
    text = aggregate_block_text(uploaded_files)
    dates_value = insights.get("Dates reperees", "Aucune")
    org_value = insights.get("Organismes reperes", "Aucun")
    amount_value = insights.get("Montants reperes", "Aucun")

    checks: list[str] = []
    score = 0

    def add_check(condition: bool, label: str) -> None:
        nonlocal score
        if condition:
            checks.append(label)
            score += 10

    if block_name == "Documents dossier":
        add_check(dates_value != "Aucune", "date de reference detectee")
        add_check(org_value != "Aucun", "organisme detecte")
        add_check(
            any(keyword in text for keyword in ["appel", "aap", "reglement", "cadre", "eligibilite", "candidature"]),
            "regles ou contexte d'appel detectes",
        )
    elif block_name == "Documents client":
        add_check(org_value != "Aucun", "identite de structure detectee")
        add_check(
            any(keyword in text for keyword in ["siret", "association", "sas", "sarl", "statuts", "adresse", "siege"]),
            "elements administratifs detectes",
        )
        add_check(
            any(keyword in text for keyword in ["activite", "presentation", "reference", "competence", "experience"]),
            "elements de presentation detectes",
        )
    elif block_name == "Documents projet":
        add_check(amount_value != "Aucun", "montant ou budget detecte")
        add_check(
            any(keyword in text for keyword in ["projet", "objectif", "action", "planning", "calendrier", "beneficiaire"]),
            "contenu projet detecte",
        )
        add_check(
            any(keyword in text for keyword in ["budget", "financement", "depense", "cout", "recette"]),
            "elements budgetaires detectes",
        )

    if score >= 30:
        status = "fort"
    elif score >= 20:
        status = "moyen"
    elif score >= 10:
        status = "faible"
    else:
        status = "insuffisant"

    return {
        "score": score,
        "status": status,
        "checks": " | ".join(checks) if checks else "aucun critere fort detecte",
    }


def collect_block_insights(uploaded_files) -> dict[str, str]:
    dates: dict[str, set[str]] = {}
    amounts: dict[str, set[str]] = {}
    organizations: dict[str, set[str]] = {}
    keywords: dict[str, set[str]] = {}
    provenances: list[str] = []

    for uploaded_file in uploaded_files:
        suffix = get_uploaded_suffix(uploaded_file)
        file_bytes = get_uploaded_bytes(uploaded_file)
        source_name = uploaded_file.name
        local_findings = []

        if suffix in {".txt", ".md"}:
            text_content = parse_text_bytes(file_bytes)
            metadata = extract_text_metadata(text_content, uploaded_file.name)
            detected_date = metadata.get("Date detectee", "")
            detected_amount = metadata.get("Montant detecte", "")
            detected_org = metadata.get("Organisme detecte", "")
            add_detected_value(dates, "date", detected_date, source_name)
            add_detected_value(amounts, "amount", detected_amount, source_name)
            add_detected_value(organizations, "org", detected_org, source_name)
            if detected_date not in {"", "Non detectee"}:
                local_findings.append(f"date {detected_date}")
            if detected_amount not in {"", "Non detecte"}:
                local_findings.append(f"montant {detected_amount}")
            if detected_org not in {"", "Non detecte"}:
                local_findings.append(f"organisme {detected_org}")
            for word in extract_keywords_from_text(text_content):
                add_detected_value(keywords, "keyword", word, source_name)

        elif suffix == ".docx":
            text_content, _, _ = parse_docx_bytes(file_bytes)
            metadata = extract_text_metadata(text_content, uploaded_file.name)
            detected_date = metadata.get("Date detectee", "")
            detected_amount = metadata.get("Montant detecte", "")
            detected_org = metadata.get("Organisme detecte", "")
            add_detected_value(dates, "date", detected_date, source_name)
            add_detected_value(amounts, "amount", detected_amount, source_name)
            add_detected_value(organizations, "org", detected_org, source_name)
            if detected_date not in {"", "Non detectee"}:
                local_findings.append(f"date {detected_date}")
            if detected_amount not in {"", "Non detecte"}:
                local_findings.append(f"montant {detected_amount}")
            if detected_org not in {"", "Non detecte"}:
                local_findings.append(f"organisme {detected_org}")
            for word in extract_keywords_from_text(text_content):
                add_detected_value(keywords, "keyword", word, source_name)

        elif suffix == ".pdf":
            pdf_text, _, _, pdf_error = parse_pdf_bytes(file_bytes)
            if pdf_error:
                provenances.append(f"{source_name} -> {pdf_error}")
                continue
            metadata = extract_text_metadata(pdf_text, uploaded_file.name)
            detected_date = metadata.get("Date detectee", "")
            detected_amount = metadata.get("Montant detecte", "")
            detected_org = metadata.get("Organisme detecte", "")
            add_detected_value(dates, "date", detected_date, source_name)
            add_detected_value(amounts, "amount", detected_amount, source_name)
            add_detected_value(organizations, "org", detected_org, source_name)
            if detected_date not in {"", "Non detectee"}:
                local_findings.append(f"date {detected_date}")
            if detected_amount not in {"", "Non detecte"}:
                local_findings.append(f"montant {detected_amount}")
            if detected_org not in {"", "Non detecte"}:
                local_findings.append(f"organisme {detected_org}")
            for word in extract_keywords_from_text(pdf_text):
                add_detected_value(keywords, "keyword", word, source_name)

        elif suffix == ".csv":
            dataframe = parse_csv_bytes(file_bytes)
            add_detected_value(keywords, "keyword", "tableau", source_name)
            for col in list(dataframe.columns)[:8]:
                add_detected_value(keywords, "keyword", str(col).lower(), source_name)
            local_findings.append(f"{len(dataframe.columns)} colonnes")
        elif suffix == ".xlsx":
            workbook = parse_excel_bytes(file_bytes)
            business_sheets, informative_sheets = filter_business_sheets(workbook)
            add_detected_value(keywords, "keyword", "tableau", source_name)
            for sheet_df in business_sheets.values():
                for col in list(sheet_df.columns)[:5]:
                    add_detected_value(keywords, "keyword", str(col).lower(), source_name)
            local_findings.append(f"{len(business_sheets)} feuille(s) metier")
            if informative_sheets:
                local_findings.append(f"{len(informative_sheets)} feuille(s) informatives")

        if local_findings:
            provenances.append(f"{source_name} -> " + ", ".join(local_findings[:4]))

    return {
        "Dates reperees": format_detected_values(dates),
        "Montants reperes": format_detected_values(amounts),
        "Organismes reperes": format_detected_values(organizations),
        "Mots-cles simples": format_detected_values(keywords),
        "Provenance synthese": " | ".join(provenances[:6]) if provenances else "Aucune",
    }


def dataframe_to_markdown(dataframe: pd.DataFrame, title: str) -> str:
    preview_df = dataframe.head(20).fillna("")
    headers = [str(col) for col in preview_df.columns]
    header_row = "| " + " | ".join(headers) + " |"
    separator_row = "| " + " | ".join(["---"] * len(headers)) + " |"

    body_rows = []
    for row in preview_df.itertuples(index=False, name=None):
        values = [str(value).replace("\n", " ").strip() for value in row]
        body_rows.append("| " + " | ".join(values) + " |")

    table_markdown = "\n".join([header_row, separator_row] + body_rows)
    return f"## {title}\n\n{table_markdown}"


def workbook_to_markdown(workbook: dict[str, pd.DataFrame], title: str) -> str:
    sections = [f"# {title}"]
    for sheet_name, sheet_df in workbook.items():
        sections.append(dataframe_to_markdown(sheet_df, sheet_name))
    return "\n\n".join(sections)


def classify_sheet(sheet_name: str) -> str:
    normalized = sheet_name.lower().strip()
    informative_keywords = [
        "lisez",
        "readme",
        "read me",
        "mode d'emploi",
        "instructions",
        "notice",
        "sommaire",
    ]
    if any(keyword in normalized for keyword in informative_keywords):
        return "informatif"
    return "metier"


def filter_business_sheets(workbook: dict[str, pd.DataFrame]) -> tuple[dict[str, pd.DataFrame], list[str]]:
    business_sheets: dict[str, pd.DataFrame] = {}
    informative_sheets: list[str] = []

    for sheet_name, sheet_df in workbook.items():
        category = classify_sheet(sheet_name)
        if category == "informatif":
            informative_sheets.append(sheet_name)
        else:
            business_sheets[sheet_name] = sheet_df

    if not business_sheets:
        return workbook, informative_sheets

    return business_sheets, informative_sheets


def render_normalized_text(content: str, filename: str) -> None:
    st.markdown("### Source normalisee")
    st.text_area("Contenu normalise", content[:5000], height=260)
    st.download_button(
        "Telecharger la source normalisee",
        data=content,
        file_name=f"{Path(filename).stem}_normalise.md",
        mime="text/markdown",
    )


def build_upload_summary(uploaded_file) -> dict[str, str]:
    suffix = get_uploaded_suffix(uploaded_file)
    return {
        "Nom": uploaded_file.name,
        "Type": suffix,
        "Taille": f"{uploaded_file.size} octets",
        "Statut": "Charge",
    }


def render_global_summary(summary_map: dict[str, list[dict[str, str]]]) -> None:
    st.subheader("Recapitulatif global")
    col1, col2, col3 = st.columns(3)

    loaded_count = sum(1 for value in summary_map.values() if value)
    col1.metric("Blocs complets", str(loaded_count))
    col2.metric("Blocs manquants", str(len(summary_map) - loaded_count))
    col3.metric("Documents charges", str(sum(len(value) for value in summary_map.values())))

    for block_name, infos in summary_map.items():
        if not infos:
            st.warning(f"{block_name} : aucun document charge")
        else:
            st.success(f"{block_name} : {len(infos)} document(s)")
            for info in infos:
                st.write(f"- {info['Nom']} | {info['Type']} | {info['Taille']}")


def build_global_cross_block_summary(block_files_map: dict[str, list]) -> dict[str, str]:
    dossier_priority = ["Documents dossier", "Documents projet", "Documents client"]
    client_priority = ["Documents client", "Documents projet", "Documents dossier"]
    project_priority = ["Documents projet", "Documents dossier", "Documents client"]

    available_blocks = [name for name, files in block_files_map.items() if files]
    missing_blocks = [name for name, files in block_files_map.items() if not files]

    block_insights = {
        name: collect_block_insights(files) if files else {}
        for name, files in block_files_map.items()
    }
    block_criteria = {
        name: evaluate_block_criteria(name, files, block_insights.get(name, {})) if files else {
            "score": 0,
            "status": "insuffisant",
            "checks": "bloc vide",
        }
        for name, files in block_files_map.items()
    }

    organizations = {}
    dates = {}
    amounts = {}
    block_statuses = {
        name: assess_block_completeness(files).get("Statut bloc", "vide")
        for name, files in block_files_map.items()
    }

    for block_name, insights in block_insights.items():
        if not insights:
            continue
        org_value = insights.get("Organismes reperes", "Aucun")
        date_value = insights.get("Dates reperees", "Aucune")
        amount_value = insights.get("Montants reperes", "Aucun")
        if org_value != "Aucun":
            organizations[block_name] = org_value
        if date_value != "Aucune":
            dates[block_name] = date_value
        if amount_value != "Aucun":
            amounts[block_name] = amount_value

    strong_blocks = [
        name for name, criteria in block_criteria.items()
        if criteria.get("status") in {"fort", "moyen"}
    ]

    if len(available_blocks) == 3 and len(strong_blocks) == 3:
        readiness = "pret pour pre-analyse"
    elif len(available_blocks) == 3 and len(strong_blocks) >= 2:
        readiness = "partiellement exploitable"
    elif len(available_blocks) == 3:
        readiness = "structure complete mais informations faibles"
    elif len(available_blocks) == 2:
        readiness = "partiellement pret"
    else:
        readiness = "insuffisant pour comparaison"

    checks = []
    issues = []
    if "Documents dossier" in organizations and "Documents projet" in organizations:
        checks.append("dossier/projet : organismes detectes disponibles")
    if "Documents client" in available_blocks and "Documents projet" in available_blocks:
        checks.append("client/projet : comparaison possible")
    if missing_blocks:
        checks.append("blocs manquants : " + ", ".join(missing_blocks))
        issues.append("au moins un bloc est manquant")

    weak_blocks = [name for name, status in block_statuses.items() if status != "suffisant"]
    if weak_blocks:
        issues.append("blocs a renforcer : " + ", ".join(weak_blocks))

    low_criteria_blocks = [
        name for name, criteria in block_criteria.items()
        if criteria.get("status") in {"faible", "insuffisant"}
    ]
    if low_criteria_blocks:
        issues.append("criteres insuffisamment identifies dans : " + ", ".join(low_criteria_blocks))

    if "Documents dossier" in dates and "Documents projet" not in dates:
        issues.append("date detectee dans le dossier mais absente du projet")
    if "Documents dossier" in organizations and "Documents client" not in organizations:
        issues.append("organisme detecte dans le dossier mais absent du bloc client")
    if "Documents projet" in amounts and "Documents dossier" not in amounts:
        issues.append("montant detecte dans le projet mais absent du dossier")

    if not organizations:
        issues.append("aucun organisme clairement detecte a l'echelle globale")
    if not dates:
        issues.append("aucune date clairement detectee a l'echelle globale")
    if not amounts:
        issues.append("aucun montant clairement detecte a l'echelle globale")

    if len(set(organizations.values())) > 1 and len(organizations) >= 2:
        issues.append("organismes detectes differents selon les blocs")
    if len(set(dates.values())) > 1 and len(dates) >= 2:
        issues.append("dates detectees non homogenes entre les blocs")
    if len(set(amounts.values())) > 1 and len(amounts) >= 2:
        issues.append("montants detectes non homogenes entre les blocs")

    if not issues:
        issues_text = "Aucune incoherence simple detectee"
    else:
        issues_text = " | ".join(issues[:8])

    prescore_label, prescore_value = compute_global_prescore(
        available_blocks=available_blocks,
        block_criteria_scores={name: int(criteria.get("score", 0)) for name, criteria in block_criteria.items()},
        issues=issues,
    )

    return {
        "Etat global": readiness,
        "Pre-score global": f"{prescore_label} ({prescore_value})",
        "Blocs disponibles": ", ".join(available_blocks) if available_blocks else "Aucun",
        "Blocs manquants": ", ".join(missing_blocks) if missing_blocks else "Aucun",
        "Statut des blocs": " | ".join(f"{k}: {v}" for k, v in block_statuses.items()),
        "Criteres dossier": f"{block_criteria['Documents dossier']['status']} ({block_criteria['Documents dossier']['score']}/30) - {block_criteria['Documents dossier']['checks']}",
        "Criteres client": f"{block_criteria['Documents client']['status']} ({block_criteria['Documents client']['score']}/30) - {block_criteria['Documents client']['checks']}",
        "Criteres projet": f"{block_criteria['Documents projet']['status']} ({block_criteria['Documents projet']['score']}/30) - {block_criteria['Documents projet']['checks']}",
        "Date prioritaire": choose_priority_value(dates, dossier_priority),
        "Organisme prioritaire": choose_priority_value(organizations, client_priority),
        "Montant prioritaire": choose_priority_value(amounts, project_priority),
        "Organismes par bloc": " | ".join(f"{k}: {v}" for k, v in organizations.items()) if organizations else "Aucun",
        "Dates par bloc": " | ".join(f"{k}: {v}" for k, v in dates.items()) if dates else "Aucune",
        "Montants par bloc": " | ".join(f"{k}: {v}" for k, v in amounts.items()) if amounts else "Aucun",
        "Controle simple": " | ".join(checks) if checks else "Aucun controle disponible",
        "Incoherences detectees": issues_text,
    }


def assess_block_completeness(uploaded_files) -> dict[str, str]:
    if not uploaded_files:
        return {
            "Statut bloc": "vide",
            "Niveau": "0/3",
            "Commentaire": "Aucun document n'a ete charge dans ce bloc.",
        }

    suffixes = {
        get_uploaded_suffix(uploaded_file)
        for uploaded_file in uploaded_files
    }
    file_count = len(uploaded_files)

    if file_count >= 2 and len(suffixes) >= 2:
        return {
            "Statut bloc": "suffisant",
            "Niveau": "3/3",
            "Commentaire": "Le bloc contient plusieurs pieces et au moins deux formats ou types apparents.",
        }

    if file_count >= 1:
        return {
            "Statut bloc": "partiel",
            "Niveau": "2/3",
            "Commentaire": "Le bloc contient deja des pieces, mais semble encore incomplet pour une analyse plus fiable.",
        }

    return {
        "Statut bloc": "vide",
        "Niveau": "0/3",
        "Commentaire": "Aucun document n'a ete charge dans ce bloc.",
    }


def render_block_summary(title: str, uploaded_files) -> None:
    st.markdown(f"## Synthese du bloc : {title}")

    type_counts: dict[str, int] = {}
    total_size = 0
    for uploaded_file in uploaded_files:
        suffix = get_uploaded_suffix(uploaded_file)
        type_counts[suffix] = type_counts.get(suffix, 0) + 1
        total_size += uploaded_file.size

    col1, col2, col3 = st.columns(3)
    col1.metric("Documents", str(len(uploaded_files)))
    col2.metric("Types detectes", str(len(type_counts)))
    col3.metric("Taille totale", f"{total_size} octets")

    st.write(
        "Repartition : "
        + ", ".join(f"`{file_type}` x {count}" for file_type, count in sorted(type_counts.items()))
    )
    render_metadata(assess_block_completeness(uploaded_files))
    render_metadata(collect_block_insights(uploaded_files))


def build_block_normalized_text(title: str, uploaded_files) -> str:
    chunks = [f"# {title}"]
    for uploaded_file in uploaded_files:
        suffix = get_uploaded_suffix(uploaded_file)
        file_bytes = get_uploaded_bytes(uploaded_file)
        chunks.append(f"## DOCUMENT: {uploaded_file.name}")
        chunks.append(f"Type: {suffix}")

        if suffix in {".txt", ".md"}:
            text_content = parse_text_bytes(file_bytes)
            chunks.append(text_content[:10000] or "[contenu vide]")
        elif suffix == ".csv":
            dataframe = parse_csv_bytes(file_bytes)
            chunks.append(dataframe_to_markdown(dataframe, uploaded_file.name))
        elif suffix == ".xlsx":
            workbook = parse_excel_bytes(file_bytes)
            business_sheets, informative_sheets = filter_business_sheets(workbook)
            if informative_sheets:
                chunks.append("## FEUILLES INFORMATIVES")
                chunks.append(", ".join(informative_sheets))
            for sheet_name, sheet_df in business_sheets.items():
                chunks.append(dataframe_to_markdown(sheet_df, f"{uploaded_file.name} - {sheet_name}"))
        elif suffix == ".pdf":
            pdf_text, _, _, pdf_error = parse_pdf_bytes(file_bytes)
            if pdf_error:
                chunks.append(f"[{pdf_error}]")
            else:
                chunks.append(pdf_text[:10000] or "[aucun texte exploitable]")
        elif suffix == ".docx":
            _, markdown_content, _ = parse_docx_bytes(file_bytes)
            chunks.append(markdown_content[:10000] or "[aucun texte exploitable]")
        else:
            chunks.append("[type non pris en charge]")

    return "\n\n".join(chunks)


def process_uploaded_file(uploaded_file, category_label: str, file_index: int) -> None:
    st.markdown(f"### {category_label} - fichier {file_index}")
    st.success("Document recu avec succes.")

    col1, col2 = st.columns(2)
    col1.write(f"Nom : `{uploaded_file.name}`")
    col2.write(f"Taille : `{uploaded_file.size}` octets")

    suffix = get_uploaded_suffix(uploaded_file)
    file_bytes = get_uploaded_bytes(uploaded_file)
    st.write(f"Type detecte : `{suffix}`")

    if suffix in {".txt", ".md"}:
        text_content = parse_text_bytes(file_bytes)
        render_metadata(extract_text_metadata(text_content, uploaded_file.name))
        st.markdown("### Apercu texte")
        st.text_area("Contenu detecte", text_content[:5000], height=250)
        render_normalized_text(text_content, uploaded_file.name)
        st.write("Etape suivante : nettoyer et structurer ce texte.")
        return

    if suffix == ".csv":
        dataframe = parse_csv_bytes(file_bytes)
        render_metadata(extract_table_metadata(dataframe, uploaded_file.name))
        st.markdown("### Apercu tabulaire")
        st.dataframe(dataframe, use_container_width=True)
        st.write(f"Nombre de lignes : `{len(dataframe)}`")
        st.write(f"Nombre de colonnes : `{len(dataframe.columns)}`")
        normalized_text = dataframe_to_markdown(dataframe, uploaded_file.name)
        render_normalized_text(normalized_text, uploaded_file.name)
        return

    if suffix == ".xlsx":
        workbook = parse_excel_bytes(file_bytes)
        sheet_names = list(workbook.keys())
        business_sheets, informative_sheets = filter_business_sheets(workbook)
        displayed_sheets = business_sheets if business_sheets else workbook

        st.markdown("### Feuilles detectees")
        st.write(", ".join(f"`{name}`" for name in sheet_names))
        if informative_sheets:
            st.info("Feuilles informatives detectees : " + ", ".join(f"`{name}`" for name in informative_sheets))

        first_sheet_name = next(iter(displayed_sheets.keys()))
        first_df = displayed_sheets[first_sheet_name]
        metadata = extract_table_metadata(first_df, uploaded_file.name)
        metadata["Nombre de feuilles"] = str(len(sheet_names))
        metadata["Feuilles metier"] = str(len(displayed_sheets))
        render_metadata(metadata)

        st.markdown("### Apercu Excel par feuille")
        for sheet_name, sheet_df in displayed_sheets.items():
            csv_content = sheet_df.to_csv(index=False)
            with st.expander(f"Feuille : {sheet_name}", expanded=False):
                st.dataframe(sheet_df, use_container_width=True)
                st.write(f"Lignes : `{len(sheet_df)}`")
                st.write(f"Colonnes : `{len(sheet_df.columns)}`")
                st.text_area(
                    f"CSV genere - {sheet_name}",
                    csv_content[:5000],
                    height=180,
                    key=f"csv_preview_{uploaded_file.name}_{sheet_name}",
                )
                st.download_button(
                    f"Telecharger {sheet_name} en CSV",
                    data=csv_content,
                    file_name=f"{Path(uploaded_file.name).stem}_{sheet_name}.csv",
                    mime="text/csv",
                    key=f"csv_download_{uploaded_file.name}_{sheet_name}",
                )
        normalized_text = workbook_to_markdown(workbook, uploaded_file.name)
        render_normalized_text(normalized_text, uploaded_file.name)
        return

    if suffix == ".pdf":
        pdf_text, page_count, text_page_count, pdf_error = parse_pdf_bytes(file_bytes)

        if pdf_error:
            st.error(f"Lecture PDF impossible : {pdf_error}")
            st.write("Ce fichier peut etre corrompu, mal exporte, ou non conforme au format PDF attendu.")
            st.write("Le bloc continue a fonctionner, mais ce document n'est pas exploitable pour l'instant.")
            return

        if not pdf_text:
            st.warning("Le PDF a ete charge, mais aucun texte exploitable n'a ete detecte.")
            st.write("Il est possible que le document soit scanne ou image.")
            return

        render_metadata(extract_text_metadata(pdf_text, uploaded_file.name))
        st.markdown("### Apercu PDF")
        st.text_area("Texte detecte", pdf_text[:5000], height=300)
        st.write(f"Pages lues : `{page_count}`")
        st.write(f"Pages avec texte : `{text_page_count}`")
        render_normalized_text(pdf_text, uploaded_file.name)
        return

    if suffix == ".docx":
        text_content, markdown_content, paragraph_count = parse_docx_bytes(file_bytes)

        if not text_content:
            st.warning("Le document DOCX a ete charge, mais aucun texte exploitable n'a ete trouve.")
            return

        render_metadata(extract_text_metadata(text_content, uploaded_file.name))
        st.markdown("### Apercu DOCX")
        st.text_area("Texte detecte", text_content[:5000], height=300)
        st.write(f"Paragraphes detectes : `{paragraph_count}`")
        st.markdown("### Conversion Markdown")
        st.text_area("Markdown genere", markdown_content[:5000], height=260)
        st.download_button(
            "Telecharger en Markdown",
            data=markdown_content,
            file_name=f"{Path(uploaded_file.name).stem}.md",
            mime="text/markdown",
        )
        render_normalized_text(markdown_content, uploaded_file.name)
        return

    st.write("Etape suivante : lecture du contenu et extraction de texte.")


def render_home() -> None:
    st.title("AAP Ingenia")
    st.caption("Prototype simple pour cadrer un futur outil d'analyse de dossiers")

    col1, col2, col3 = st.columns(3)
    col1.metric("Statut", "Prototype en ligne")
    col2.metric("Mode", "Simple et stable")
    col3.metric("Donnees demo", "Disponibles")

    st.markdown(
        """
        Cette version sert a garder une base propre et facile a deployer.

        Elle permet deja de :
        - presenter le projet ;
        - visualiser des donnees de demonstration ;
        - tester un premier upload de document ;
        - preparer la suite sans complexite inutile.
        """
    )


def render_project() -> None:
    st.subheader("Ou en est le projet ?")
    st.markdown(
        """
        - Le cadrage produit est deja bien avance
        - Le schema de donnees a ete reflechi
        - L'application web commence simplement
        - La connexion a une base reelle viendra plus tard
        """
    )

    st.subheader("Prochaine etape conseillee")
    st.write(
        "Construire un parcours d'upload simple, puis afficher les metadonnees du fichier avant toute analyse automatisee."
    )


def render_demo_data() -> None:
    st.subheader("Donnees de demonstration")
    df = load_demo_data()

    if df.empty:
        st.warning("Le fichier de demonstration est introuvable.")
        return

    st.dataframe(df, use_container_width=True)

    first_row = df.iloc[0]
    st.markdown("### Resume")
    st.write(first_row.get("resume_executif", "Aucun resume disponible."))


def render_upload_block(title: str, help_text: str, uploader_key: str):
    st.subheader(title)
    st.caption(help_text)
    uploaded_files = st.file_uploader(
        f"Depose un ou plusieurs documents pour : {title}",
        type=["pdf", "docx", "txt", "md", "csv", "xlsx"],
        key=uploader_key,
        accept_multiple_files=True,
    )

    if not uploaded_files:
        st.info("Aucun document charge pour ce bloc.")
        return []

    for index, uploaded_file in enumerate(uploaded_files, start=1):
        process_uploaded_file(uploaded_file, title, index)
        if index < len(uploaded_files):
            st.divider()

    st.divider()
    render_block_summary(title, uploaded_files)
    block_normalized_text = build_block_normalized_text(title, uploaded_files)
    render_normalized_text(block_normalized_text, title.replace(" ", "_").lower())

    return uploaded_files


def render_upload() -> None:
    st.subheader("Upload structure en 3 blocs")
    st.write(
        "Le flux metier repose sur trois types de documents distincts : dossier, client et projet."
    )

    summary_map = {}
    block_files_map = {}

    block_files_map["Documents dossier"] = render_upload_block(
        "Documents dossier",
        "Documents cibles a analyser : appel a projets, reglement, cahier des charges, cadre d'intervention.",
        "upload_dossier",
    )
    summary_map["Documents dossier"] = [
        build_upload_summary(uploaded_file) for uploaded_file in block_files_map["Documents dossier"]
    ]
    st.divider()
    block_files_map["Documents client"] = render_upload_block(
        "Documents client",
        "Documents qui decrivent la structure porteuse : presentation, statuts, references, plaquette.",
        "upload_client",
    )
    summary_map["Documents client"] = [
        build_upload_summary(uploaded_file) for uploaded_file in block_files_map["Documents client"]
    ]
    st.divider()
    block_files_map["Documents projet"] = render_upload_block(
        "Documents projet",
        "Documents qui decrivent l'action ou la demande : note d'intention, budget, description du projet.",
        "upload_projet",
    )
    summary_map["Documents projet"] = [
        build_upload_summary(uploaded_file) for uploaded_file in block_files_map["Documents projet"]
    ]
    st.divider()
    render_global_summary(summary_map)
    st.divider()
    st.subheader("Synthese globale inter-blocs")
    render_metadata(build_global_cross_block_summary(block_files_map))


def main() -> None:
    st.set_page_config(
        page_title="AAP Ingenia",
        page_icon="📁",
        layout="wide",
    )

    page = st.sidebar.radio(
        "Navigation",
        [
            "Accueil",
            "Projet",
            "Donnees demo",
            "Upload",
        ],
    )

    if page == "Accueil":
        render_home()
    elif page == "Projet":
        render_project()
    elif page == "Donnees demo":
        render_demo_data()
    else:
        render_upload()


if __name__ == "__main__":
    main()
