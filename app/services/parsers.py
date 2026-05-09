from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pandas as pd
from docx import Document
from pypdf import PdfReader
from pypdf.errors import PdfReadError


def _table_to_markdown(table) -> str:
    """Convertit un tableau python-docx en markdown (| col | col |)."""
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            # Ligne de séparation header
            rows.append("| " + " | ".join("---" for _ in cells) + " |")
    return "\n".join(rows)


def docx_to_markdown(document: Document) -> str:
    """Convertit un Document Word en Markdown.

    Traite dans l'ordre du document :
    - Paragraphes (texte + titres)
    - Tableaux (convertis en tableau Markdown)
    """
    from docx.oxml.ns import qn  # noqa: PLC0415

    blocks = []

    # Itération sur les éléments body dans l'ordre réel du document
    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            # Paragraphe
            para_text = "".join(
                node.text or "" for node in child.iter(qn("w:t"))
            ).strip()
            if not para_text:
                continue
            style_name = (child.find(f".//{qn('w:pStyle')}") or child).get(
                qn("w:val"), ""
            ) or ""
            if "heading" in style_name.lower() or "titre" in style_name.lower():
                blocks.append(f"## {para_text}")
            else:
                blocks.append(para_text)

        elif tag == "tbl":
            # Tableau — reconstruit via python-docx Table object
            try:
                from docx.table import Table  # noqa: PLC0415
                tbl = Table(child, document)
                md = _table_to_markdown(tbl)
                if md.strip():
                    blocks.append(md)
            except Exception:
                pass  # tableau non parsable, on l'ignore silencieusement

    return "\n\n".join(blocks)


def get_uploaded_suffix(uploaded_file) -> str:
    return Path(uploaded_file.name).suffix.lower() or "inconnu"


def get_uploaded_bytes(uploaded_file) -> bytes:
    return uploaded_file.getvalue()


def parse_text_bytes(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")


def parse_csv_bytes(file_bytes: bytes) -> pd.DataFrame:
    return pd.read_csv(BytesIO(file_bytes))


def parse_excel_bytes(file_bytes: bytes) -> dict[str, pd.DataFrame]:
    return pd.read_excel(BytesIO(file_bytes), sheet_name=None)


def parse_pdf_bytes(file_bytes: bytes) -> tuple[str, int, int, str | None]:
    try:
        reader = PdfReader(BytesIO(file_bytes))
        pages_text = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages_text.append(page_text.strip())
        return "\n\n".join(pages_text), len(reader.pages), len(pages_text), None
    except PdfReadError as exc:
        return "", 0, 0, f"PDF illisible par le parseur ({exc})"
    except Exception as exc:
        return "", 0, 0, f"Erreur de lecture PDF ({exc.__class__.__name__})"


def parse_docx_bytes(file_bytes: bytes) -> tuple[str, str, int]:
    document = Document(BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    # text_content inclut aussi le texte des cellules de tableaux
    table_texts = []
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                table_texts.append(row_text)
    all_texts = paragraphs + table_texts
    text_content = "\n\n".join(all_texts)
    markdown_content = docx_to_markdown(document)
    return text_content, markdown_content, len(all_texts)
